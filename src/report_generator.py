#!/usr/bin/env python3
"""
Report Generation Module
Generates attendance and emotion reports in TXT/DOCX format
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from config import *

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generate attendance reports"""
    
    def __init__(self):
        pass
    
    def generate_report(self, attendance, emotion_summary, subject, time_start, time_end, report_format='both', image_path=None):
        """
        Generate attendance report
        
        Args:
            attendance: Dict of student attendance {name: status}
            emotion_summary: Dict of emotion percentages
            subject: Subject name
            time_start: Start time string
            time_end: End time string
            report_format: 'txt', 'docx', or 'both'
            image_path: Optional path to captured class image to embed
            
        Returns:
            list: Paths to generated report files
        """
        report_files = []
        
        try:
            # Generate filename
            timestamp = datetime.now().strftime(REPORT_TIMESTAMP_FORMAT)
            base_filename = f"{subject}_{timestamp}"
            
            # Generate TXT report
            if report_format in ['txt', 'both']:
                txt_file = self._generate_txt_report(
                    attendance, emotion_summary, subject,
                    time_start, time_end, base_filename
                )
                if txt_file:
                    report_files.append(txt_file)
            
            # Generate DOCX report
            if report_format in ['docx', 'both']:
                docx_file = self._generate_docx_report(
                    attendance, emotion_summary, subject,
                    time_start, time_end, base_filename, image_path
                )
                if docx_file:
                    report_files.append(docx_file)
            
            logger.info(f"Generated {len(report_files)} report(s)")
            return report_files
            
        except Exception as e:
            logger.error(f"Error generating reports: {e}")
            return []
    
    def _generate_txt_report(self, attendance, emotion_summary, subject, time_start, time_end, base_filename):
        """Generate TXT format report"""
        try:
            filepath = os.path.join(REPORTS_DIR, f"{base_filename}.txt")
            
            # Prepare data
            date_str = datetime.now().strftime("%Y-%m-%d")
            present = [name for name, status in attendance.items() if status == "Present"]
            absent = [name for name, status in attendance.items() if status == "Absent"]
            total = len(attendance)
            attendance_rate = (len(present) / total * 100) if total > 0 else 0
            
            # Write report
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 70 + "\n")
                f.write(f"ATTENDANCE REPORT\n")
                f.write("=" * 70 + "\n\n")
                
                f.write(f"Subject: {subject}\n")
                f.write(f"Date: {date_str}\n")
                f.write(f"Time: {time_start} – {time_end}\n\n")
                
                f.write("-" * 70 + "\n")
                f.write("ATTENDANCE SUMMARY\n")
                f.write("-" * 70 + "\n\n")
                
                f.write(f"Total Students: {total}\n")
                f.write(f"Present: {len(present)}\n")
                f.write(f"Absent: {len(absent)}\n")
                f.write(f"Attendance Rate: {attendance_rate:.1f}%\n\n")
                
                f.write("-" * 70 + "\n")
                f.write("PRESENT STUDENTS\n")
                f.write("-" * 70 + "\n\n")
                
                if present:
                    for i, name in enumerate(sorted(present), 1):
                        f.write(f"{i}. {name}\n")
                else:
                    f.write("No students present\n")
                
                f.write("\n")
                f.write("-" * 70 + "\n")
                f.write("ABSENT STUDENTS\n")
                f.write("-" * 70 + "\n\n")
                
                if absent:
                    for i, name in enumerate(sorted(absent), 1):
                        f.write(f"{i}. {name}\n")
                else:
                    f.write("No students absent\n")
                
                if emotion_summary:
                    f.write("\n")
                    f.write("-" * 70 + "\n")
                    f.write("CLASS EMOTION SUMMARY\n")
                    f.write("-" * 70 + "\n\n")
                    
                    for emotion, percentage in emotion_summary.items():
                        bar = "█" * int(percentage / 5)
                        f.write(f"{emotion:12} {percentage:5.1f}%  {bar}\n")
                
                f.write("\n")
                f.write("=" * 70 + "\n")
                f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("Attendance & Emotion Analytics System\n")
                f.write("=" * 70 + "\n")
            
            logger.info(f"TXT report generated: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error generating TXT report: {e}")
            return None
    
    def _generate_docx_report(self, attendance, emotion_summary, subject, time_start, time_end, base_filename, image_path=None):
        """Generate DOCX format report"""
        try:
            filepath = os.path.join(REPORTS_DIR, f"{base_filename}.docx")
            
            # Create document
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Title
            title = doc.add_heading('Attendance Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            date_str = datetime.now().strftime("%Y-%m-%d")
            metadata = doc.add_paragraph()
            metadata.add_run(f"Subject: ").bold = True
            metadata.add_run(f"{subject}\n")
            metadata.add_run(f"Date: ").bold = True
            metadata.add_run(f"{date_str}\n")
            metadata.add_run(f"Time: ").bold = True
            metadata.add_run(f"{time_start} – {time_end}")
            
            # Class Photo
            if image_path and os.path.exists(image_path):
                doc.add_heading('Class Capture', 1)
                try:
                    doc.add_picture(image_path, width=Inches(6.0))
                except Exception as e:
                    logger.warning(f"Could not add picture to report: {e}")
            
            # Summary
            doc.add_heading('Attendance Summary', 1)
            
            present = [name for name, status in attendance.items() if status == "Present"]
            absent = [name for name, status in attendance.items() if status == "Absent"]
            total = len(attendance)
            attendance_rate = (len(present) / total * 100) if total > 0 else 0
            
            # Summary table
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            table.rows[0].cells[0].text = 'Total Students'
            table.rows[0].cells[1].text = str(total)
            table.rows[1].cells[0].text = 'Present'
            table.rows[1].cells[1].text = str(len(present))
            table.rows[2].cells[0].text = 'Absent'
            table.rows[2].cells[1].text = str(len(absent))
            table.rows[3].cells[0].text = 'Attendance Rate'
            table.rows[3].cells[1].text = f"{attendance_rate:.1f}%"
            
            # Present students
            doc.add_heading('Present Students', 1)
            if present:
                for i, name in enumerate(sorted(present), 1):
                    doc.add_paragraph(f"{i}. {name}", style='List Number')
            else:
                doc.add_paragraph("No students present")
            
            # Absent students
            doc.add_heading('Absent Students', 1)
            if absent:
                for i, name in enumerate(sorted(absent), 1):
                    doc.add_paragraph(f"{i}. {name}", style='List Number')
            else:
                doc.add_paragraph("No students absent")
            
            # Emotion summary
            if emotion_summary:
                doc.add_heading('Class Emotion Summary', 1)
                
                for emotion, percentage in emotion_summary.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{emotion}: ").bold = True
                    p.add_run(f"{percentage:.1f}%")
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
            footer.add_run("Attendance & Emotion Analytics System").italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(filepath)
            logger.info(f"DOCX report generated: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}")
            return None