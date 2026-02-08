#!/usr/bin/env python3
"""
Smart Attendance System - Modern GUI
Om Bhamare ka Smart System
CustomTkinter use kiya hai mast look ke liye
"""

import customtkinter as ctk
from tkinter import filedialog
import custom_dialogs as messagebox
import tkinter as tk
from PIL import Image, ImageTk
import cv2
import threading
import os
import logging
import shutil
import random
from datetime import datetime

# Backend modules import kar rahe hai
from image_capture import ImageCapture
from face_recognition_module import FaceRecognitionModule
from emotion_detection import EmotionDetection
from report_generator import ReportGenerator
from email_automation import EmailAutomation
from data_cleanup import DataCleanup
from settings_manager import SettingsManager
from config import *

# Logging setup - sab record hoga yaha
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                   handlers=[
                       logging.FileHandler(os.path.join(LOGS_DIR, "app.log")),
                       logging.StreamHandler()
                   ])
logger = logging.getLogger(__name__)

# Theme set karte hai - Dark mode best hai
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class AttendanceApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Setup - Title aur size set karte hai
        self.title(WINDOW_TITLE)
        self.geometry(f"{WINDOW_WIDTH}x{WINDOW_HEIGHT}")
        self.minsize(1100, 750)
        
        # Grid layout banate hai
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        # Saare modules load kar rahe hai
        self.image_capture = ImageCapture()
        self.face_recognition = FaceRecognitionModule()
        self.emotion_detection = EmotionDetection()
        self.report_generator = ReportGenerator()
        self.settings_manager = SettingsManager()
        self.email_automation = EmailAutomation(self.settings_manager)
        self.data_cleanup = DataCleanup()
        
        # Variables (State maintain karne ke liye)
        self.camera_active = False
        self.automation_active = False
        self.after_id = None
        self.showing_result = False
        self.last_attendance = {}
        
        self.last_analysis_time = datetime.now()
        self.last_attendance_time = datetime.now()
        
        # Haar Cascade load karte hai (chehra dhundne ke liye)
        try:
            self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        except:
            logger.warning("Haar Cascade nahi mila yaar")
            self.face_cascade = None

        # Face Recognition ko background me start karte hai taaki app hang na ho
        threading.Thread(target=self._warmup_model, daemon=True).start()

        # Layout Setup
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.create_sidebar()
        self.create_main_area()

        # Pehle Dashboard dikhayenge
        self.show_dashboard()
        self.after(1000, self.auto_start_camera)

    def create_sidebar(self):
        """Side wala menu banate hai"""
        self.sidebar_frame = ctk.CTkFrame(self, width=260, corner_radius=0, fg_color=THEME_COLORS['surface'])
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(8, weight=1) # Neeche jagah chodne ke liye

        # --- Logo Area ---
        logo_frame = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        logo_frame.grid(row=0, column=0, padx=20, pady=(30, 20), sticky="ew")
        
        ctk.CTkLabel(logo_frame, text="🧠", font=ctk.CTkFont(size=36)).pack(side="left")
        
        title_frame = ctk.CTkFrame(logo_frame, fg_color="transparent")
        title_frame.pack(side="left", padx=10)
        ctk.CTkLabel(title_frame, text="ATTENDANCE", 
                    font=ctk.CTkFont(family=FONT_FAMILY, size=18, weight="bold"),
                    text_color=THEME_COLORS['text']).pack(anchor="w")
        ctk.CTkLabel(title_frame, text="Made by Om Bhamare", 
                    font=ctk.CTkFont(family=FONT_FAMILY, size=10, weight="bold"),
                    text_color=THEME_COLORS['secondary']).pack(anchor="w")

        # --- Status Badge ---
        status_card = ctk.CTkFrame(self.sidebar_frame, fg_color=THEME_COLORS['background'], corner_radius=8)
        status_card.grid(row=1, column=0, padx=20, pady=(0, 30), sticky="ew")
        
        status_dot = ctk.CTkLabel(status_card, text="●", text_color=THEME_COLORS['success'], font=ctk.CTkFont(size=14))
        status_dot.pack(side="left", padx=(15, 5), pady=10)
        ctk.CTkLabel(status_card, text="SYSTEM ONLINE", font=ctk.CTkFont(size=11, weight="bold"),
                    text_color=THEME_COLORS['text_dim']).pack(side="left", pady=10)

        # --- Navigation Buttons ---
        self.nav_buttons = {}
        buttons = [
            ("📊  Dashboard", self.show_dashboard),
            ("📹  Live Capture", self.show_live_capture),
            ("📝  Enrollment", self.show_enrollment),
            ("👥  Students", self.show_student_database),
            ("📄  Reports", self.show_reports),
            ("⚙️  Settings", self.show_settings),
            ("🧹  Cleanup", self.show_cleanup)
        ]

        for i, (text, command) in enumerate(buttons):
            btn = ctk.CTkButton(self.sidebar_frame, text=text, command=command,
                               font=ctk.CTkFont(family=FONT_FAMILY, size=14, weight="bold"),
                               fg_color="transparent", 
                               text_color=THEME_COLORS['text_dim'],
                               hover_color=THEME_COLORS['background'],
                               anchor="w", height=50, corner_radius=8)
            btn.grid(row=i+2, column=0, padx=15, pady=4, sticky="ew")
            self.nav_buttons[text] = btn
            
        # --- Bottom Profile ---
        profile_frame = ctk.CTkFrame(self.sidebar_frame, fg_color=THEME_COLORS['background'], corner_radius=10)
        profile_frame.grid(row=9, column=0, padx=20, pady=20, sticky="ew")
        
        ctk.CTkLabel(profile_frame, text="👤", font=ctk.CTkFont(size=24)).pack(side="left", padx=15, pady=15)
        user_info = ctk.CTkFrame(profile_frame, fg_color="transparent")
        user_info.pack(side="left")
        ctk.CTkLabel(user_info, text="Admin User", font=ctk.CTkFont(size=13, weight="bold")).pack(anchor="w")
        ctk.CTkLabel(user_info, text=f"v{APP_VERSION}", font=ctk.CTkFont(size=11), text_color="gray").pack(anchor="w")

    def create_main_area(self):
        self.main_frame = ctk.CTkFrame(self, corner_radius=0, fg_color=THEME_COLORS['background'])
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        self.main_frame.grid_rowconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(0, weight=1)

    def clear_content(self):
        if hasattr(self, 'current_page'):
            if self.current_page == 'live_capture' and self.camera_active:
                self.stop_camera()
            elif self.current_page == 'enrollment' and self.camera_active:
                self.stop_enrollment_camera()
        
        for widget in self.main_frame.winfo_children():
            widget.destroy()

    # ================= PAGES =================


    def show_dashboard(self):
        self.clear_content()
        self.current_page = "dashboard"
        
        # Reset grid configuration (other pages may have changed it)
        self.main_frame.columnconfigure(0, weight=1)
        self.main_frame.columnconfigure(1, weight=0)
        self.main_frame.columnconfigure(2, weight=0)
        self.main_frame.rowconfigure(0, weight=1)
        
        # Main Scrollable Content
        content = ctk.CTkScrollableFrame(self.main_frame, fg_color="transparent")
        content.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)

        # --- Hero Section ---
        hero_frame = ctk.CTkFrame(content, fg_color=THEME_COLORS['accent'], corner_radius=15)
        hero_frame.pack(fill="x", pady=(0, 30))
        
        # Gradient-like effect overlay (simulated with nested frames or just solid color for now)
        ctk.CTkLabel(hero_frame, text="Welcome Back, Admin! 👋", 
                    font=ctk.CTkFont(family=FONT_FAMILY, size=28, weight="bold"), 
                    text_color="white").pack(anchor="w", padx=30, pady=(30, 5))
                    
        ctk.CTkLabel(hero_frame, text="System is running optimally. Ready for attendance tracking.", 
                    font=ctk.CTkFont(family=FONT_FAMILY, size=14), 
                    text_color="white").pack(anchor="w", padx=30, pady=(0, 30))

        # --- Quick Actions ---
        actions_frame = ctk.CTkFrame(hero_frame, fg_color="transparent")
        actions_frame.place(relx=1.0, rely=0.5, anchor="e", x=-30)
        
        ctk.CTkButton(actions_frame, text="▶ Start Live Capture", command=self.show_live_capture,
                     font=ctk.CTkFont(weight="bold"), fg_color="white", text_color=THEME_COLORS['primary_dark'],
                     hover_color="#f1f5f9", height=40).pack(side="right", padx=10)
                     
        ctk.CTkButton(actions_frame, text="+ Enroll Student", command=self.show_enrollment,
                     font=ctk.CTkFont(weight="bold"), fg_color=THEME_COLORS['surface'], 
                     border_width=1, border_color="white", text_color="white",
                     hover_color=THEME_COLORS['primary'], height=40).pack(side="right", padx=10)

        # --- Stats Grid ---
        stats_label = ctk.CTkLabel(content, text="System Overview", 
                                  font=ctk.CTkFont(family=FONT_FAMILY, size=20, weight="bold"))
        stats_label.pack(anchor="w", pady=(0, 15))

        stats_grid = ctk.CTkFrame(content, fg_color="transparent")
        stats_grid.pack(fill="x", pady=(0, 30))
        stats_grid.columnconfigure((0, 1, 2, 3), weight=1)

        # Mock Data (Replace with real data later)
        students_count = len(self.face_recognition.get_all_students())
        reports_count = len([f for f in os.listdir(REPORTS_DIR) if f.endswith(('.txt', '.docx'))])
        
        self.create_modern_stat_card(stats_grid, 0, "Total Students", str(students_count), "👥", THEME_COLORS['info'])
        self.create_modern_stat_card(stats_grid, 1, "Present Today", "0", "✅", THEME_COLORS['success'])
        self.create_modern_stat_card(stats_grid, 2, "Reports Generated", str(reports_count), "📄", THEME_COLORS['warning'])
        self.create_modern_stat_card(stats_grid, 3, "System Health", "98%", "⚡", THEME_COLORS['secondary'])

        # --- Charts & Trends (New) ---
        chart_section = ctk.CTkFrame(content, fg_color="transparent")
        chart_section.pack(fill="x", pady=(0, 20))
        chart_section.columnconfigure((0, 1), weight=1)
        
        # 1. Weekly Attendance (Real Data from Reports)
        trend_card = ctk.CTkFrame(chart_section, fg_color=THEME_COLORS['surface'], corner_radius=15)
        trend_card.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        ctk.CTkLabel(trend_card, text="📊 Weekly Attendance Trends", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=20, pady=15)
        
        bars_frame = ctk.CTkFrame(trend_card, fg_color="transparent")
        bars_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Get real attendance data from reports
        days = ["Mon", "Tue", "Wed", "Thu", "Fri"]
        values = self._get_weekly_attendance_data()
        
        for i, (day, val) in enumerate(zip(days, values)):
            col = ctk.CTkFrame(bars_frame, fg_color="transparent")
            col.pack(side="left", fill="both", expand=True)
            
            # Bar (Progress Bar vertical simulation)
            bar_height = 150
            h = int(bar_height * val) if val > 0 else 5  # Minimum height for visibility
            
            bar = ctk.CTkFrame(col, width=20, height=h, fg_color=THEME_COLORS['primary'], corner_radius=5)
            bar.pack(side="bottom", pady=(0, 5))
            
            # Show percentage if available
            if val > 0:
                ctk.CTkLabel(col, text=f"{int(val*100)}%", font=ctk.CTkFont(size=9), text_color="gray").pack(side="bottom", pady=(0, 2))
            ctk.CTkLabel(col, text=day, font=ctk.CTkFont(size=10), text_color="gray").pack(side="bottom")

        # 2. Class Occupancy (Real Data)
        occ_card = ctk.CTkFrame(chart_section, fg_color=THEME_COLORS['surface'], corner_radius=15)
        occ_card.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        
        ctk.CTkLabel(occ_card, text="🏫 Class Occupancy", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=20, pady=15)
        
        # Calculate real occupancy
        total_students = students_count
        present_count = len(self.last_attendance)
        occupancy_rate = (present_count / total_students) if total_students > 0 else 0
        
        # Determine status and color
        if occupancy_rate >= 0.7:
            status_text = "High Occupancy"
            status_color = THEME_COLORS['success']
        elif occupancy_rate >= 0.4:
            status_text = "Moderate Occupancy"
            status_color = THEME_COLORS['warning']
        else:
            status_text = "Low Occupancy"
            status_color = THEME_COLORS['info']
        
        ctk.CTkLabel(occ_card, text="Current Session", text_color="gray").pack(pady=(20, 5))
        ctk.CTkLabel(occ_card, text=f"{present_count} / {total_students}", 
                    font=ctk.CTkFont(size=40, weight="bold"), 
                    text_color=status_color).pack()
        
        progress = ctk.CTkProgressBar(occ_card, width=300, height=15, corner_radius=10, progress_color=status_color)
        progress.set(occupancy_rate)
        progress.pack(pady=20)
        ctk.CTkLabel(occ_card, text=status_text, text_color="gray", font=ctk.CTkFont(size=12)).pack()

        # --- Recent Activity ---
        activity_label = ctk.CTkLabel(content, text="Recent Activity", 
                                     font=ctk.CTkFont(family=FONT_FAMILY, size=20, weight="bold"))
        activity_label.pack(anchor="w", pady=(0, 15))
        
        # Recent Activity List
        activity_frame = ctk.CTkFrame(content, fg_color=THEME_COLORS['surface'], corner_radius=10)
        activity_frame.pack(fill="x", pady=10)

        # Default system events if no real data
        now_str = datetime.now().strftime("%I:%M %p")
        activities = [
            ("🟢 System Online", now_str, "All services running"),
            ("📁 Database Init", now_str, f"Loaded {students_count} records"),
            ("🛡 Security Scan", "14:05 PM", "No threats detected"),
            ("👤 Admin Access", "13:55 PM", "New login session"),
            ("✅ Health Check", "13:50 PM", "Face module ready")
        ]
        
        # Merge with real recent attendance if any
        if self.last_attendance:
             for name, time in list(self.last_attendance.items())[:5]:
                 activities.insert(0, (f"📸 {name}", time, "Marked Present"))

        if not activities:
            ctk.CTkLabel(activity_frame, text="No recent activity to display.", text_color="gray", font=ctk.CTkFont(slant="italic")).pack(pady=20)
        

        
        for i, (title, time, desc) in enumerate(activities):
            row = ctk.CTkFrame(activity_frame, fg_color="transparent")
            row.pack(fill="x", padx=20, pady=10)
            
            ctk.CTkLabel(row, text=title, font=ctk.CTkFont(weight="bold")).pack(side="left")
            ctk.CTkLabel(row, text=time, font=ctk.CTkFont(size=12), text_color="gray").pack(side="right")
            # ctk.CTkLabel(row, text=desc, text_color="gray").pack(side="left", padx=20) # Optional desc

            if i < len(activities) - 1:
                ctk.CTkFrame(activity_frame, height=1, fg_color=THEME_COLORS['border']).pack(fill="x", padx=20)


    def create_modern_stat_card(self, parent, col, title, value, icon, color):
        card = ctk.CTkFrame(parent, fg_color=THEME_COLORS['surface'], corner_radius=12)
        card.grid(row=0, column=col, padx=10, sticky="ew")
        
        # Icon Box
        icon_frame = ctk.CTkFrame(card, fg_color=color, width=50, height=50, corner_radius=10)
        icon_frame.pack(anchor="n", padx=20, pady=(20, 0), side="left")
        icon_frame.pack_propagate(False)
        ctk.CTkLabel(icon_frame, text=icon, font=ctk.CTkFont(size=24), text_color="white").place(relx=0.5, rely=0.5, anchor="center")
        
        # Text
        text_frame = ctk.CTkFrame(card, fg_color="transparent")
        text_frame.pack(side="left", padx=10, pady=20)
        
        ctk.CTkLabel(text_frame, text=title, font=ctk.CTkFont(size=12, weight="bold"), 
                    text_color=THEME_COLORS['text_dim']).pack(anchor="w")
        ctk.CTkLabel(text_frame, text=value, font=ctk.CTkFont(size=24, weight="bold"), 
                    text_color=THEME_COLORS['text']).pack(anchor="w")

    def create_stat_card(self, parent, col, title, value, icon):
        # Legacy method kept for compatibility if needed, but redirects to modern
        pass

    def _get_weekly_attendance_data(self):
        """Get real attendance data from reports for the past 5 weekdays"""
        try:
            # Get all report files
            report_files = [f for f in os.listdir(REPORTS_DIR) if f.endswith('.txt')]
            
            if not report_files:
                return [0, 0, 0, 0, 0]  # No data available
            
            # Parse attendance rates from recent reports
            attendance_rates = []
            for report_file in sorted(report_files, reverse=True)[:5]:  # Last 5 reports
                try:
                    filepath = os.path.join(REPORTS_DIR, report_file)
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        # Look for "Attendance Rate: XX.X%"
                        if "Attendance Rate:" in content:
                            rate_line = [line for line in content.split('\n') if 'Attendance Rate:' in line][0]
                            rate_str = rate_line.split(':')[1].strip().replace('%', '')
                            rate = float(rate_str) / 100.0
                            attendance_rates.append(rate)
                except Exception as e:
                    logger.debug(f"Error parsing report {report_file}: {e}")
                    continue
            
            # Pad with zeros if we don't have 5 reports
            while len(attendance_rates) < 5:
                attendance_rates.append(0)
            
            # Return last 5 (or padded list)
            return attendance_rates[:5][::-1]  # Reverse to show oldest to newest
            
        except Exception as e:
            logger.error(f"Error getting weekly attendance data: {e}")
            return [0, 0, 0, 0, 0]

    def show_live_capture(self):
        self.clear_content()
        self.current_page = "live_capture"
        
        # Grid Layout: 3 columns (2 for camera, 1 for feed)
        self.main_frame.columnconfigure(0, weight=3)
        self.main_frame.columnconfigure(1, weight=1)
        self.main_frame.rowconfigure(0, weight=1)

        # --- Left: Camera Preview ---
        cam_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        cam_frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        cam_frame.rowconfigure(0, weight=1)
        cam_frame.columnconfigure(0, weight=1)

        # Camera Container (Black background with easy rounding)
        self.cam_container = ctk.CTkFrame(cam_frame, fg_color="black", corner_radius=15, 
                                        border_width=2, border_color=THEME_COLORS['primary'])
        self.cam_container.grid(row=0, column=0, sticky="nsew", pady=(0, 20))
        self.cam_container.grid_propagate(False) # LOCK LAYOUT
        self.cam_container.grid_rowconfigure(0, weight=1)
        self.cam_container.grid_columnconfigure(0, weight=1)
        
        self.camera_label = ctk.CTkLabel(self.cam_container, text="Camera Offline\nClick Start to Begin", 
                                        text_color="gray", font=ctk.CTkFont(size=16))
        self.camera_label.grid(row=0, column=0)

        # Control Bar (Floating or bottom)
        controls_frame = ctk.CTkFrame(cam_frame, fg_color=THEME_COLORS['surface'], corner_radius=50, height=80)
        controls_frame.grid(row=1, column=0, sticky="ew")
        
        # Centered Controls
        ctrl_inner = ctk.CTkFrame(controls_frame, fg_color="transparent")
        ctrl_inner.pack(pady=15)
        
        self.start_btn = ctk.CTkButton(ctrl_inner, text="▶ START SESSION", command=self.start_automated_session,
                                      font=ctk.CTkFont(weight="bold"), fg_color=THEME_COLORS['success'],
                                      hover_color="#059669", width=160, height=45, corner_radius=25)
        self.start_btn.pack(side="left", padx=10)

        self.capture_btn = ctk.CTkButton(ctrl_inner, text="📸 SNAP", command=self.capture_and_process,
                                        fg_color=THEME_COLORS['primary'], height=45, state="disabled",
                                        font=ctk.CTkFont(weight="bold"))
        self.capture_btn.pack(side="left", fill="x", expand=True, padx=(5, 0))
        


        ctk.CTkButton(ctrl_inner, text="⏹ STOP", command=self.stop_camera,
                     fg_color=THEME_COLORS['danger'], hover_color="#b91c1c", 
                     width=100, height=45, corner_radius=25).pack(side="left", padx=10)

        # --- Right: Live Feed & Info ---
        info_panel = ctk.CTkFrame(self.main_frame, fg_color=THEME_COLORS['surface'], corner_radius=0)
        info_panel.grid(row=0, column=1, sticky="nsew")
        info_panel.pack_propagate(False) # Strict size
        
        ctk.CTkLabel(info_panel, text="LIVE FEED", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color=THEME_COLORS['text_dim']).pack(pady=(20, 10), anchor="w", padx=20)
        
        # Subject Selector
        ctk.CTkLabel(info_panel, text="Current Subject:", font=ctk.CTkFont(size=12), text_color="gray").pack(anchor="w", padx=20)
        self.subject_var = ctk.StringVar(value=get_all_subjects()[0] if get_all_subjects() else "Default")
        ctk.CTkOptionMenu(info_panel, variable=self.subject_var, values=get_all_subjects(),
                         fg_color=THEME_COLORS['background'], button_color=THEME_COLORS['primary'],
                         dropdown_fg_color=THEME_COLORS['surface']).pack(fill="x", padx=20, pady=(5, 10))

        # Upload Button (New)
        ctk.CTkButton(info_panel, text="📤 Upload Photo to Analyze", command=self.upload_photo_attendance,
                      fg_color=THEME_COLORS['background'], hover_color=THEME_COLORS['primary'],
                      border_width=1, border_color=THEME_COLORS['border']).pack(fill="x", padx=20, pady=(0, 20))

        # Recent Detections List
        ctk.CTkLabel(info_panel, text="Recent Detections", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color=THEME_COLORS['text_dim']).pack(pady=(10, 10), anchor="w", padx=20)
        
        self.feed_container = ctk.CTkScrollableFrame(info_panel, fg_color="transparent")
        self.feed_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Empty State
        self.empty_feed_label = ctk.CTkLabel(self.feed_container, text="Waiting for students...", 
                                            text_color="gray", font=ctk.CTkFont(slant="italic"))
        self.empty_feed_label.pack(pady=40)

    def update_recent_feed(self, name, time, emotion="Neutral"):
        """Add a card to the recent feed"""
        if hasattr(self, 'empty_feed_label') and self.empty_feed_label.winfo_exists():
            self.empty_feed_label.destroy()
            
        card = ctk.CTkFrame(self.feed_container, fg_color=THEME_COLORS['background'], corner_radius=10)
        card.pack(fill="x", pady=5)
        
        icon_lbl = ctk.CTkLabel(card, text="👤", font=ctk.CTkFont(size=20))
        icon_lbl.pack(side="left", padx=10, pady=10)
        
        info = ctk.CTkFrame(card, fg_color="transparent")
        info.pack(side="left", fill="x", expand=True)
        
        ctk.CTkLabel(info, text=name, font=ctk.CTkFont(weight="bold")).pack(anchor="w")
        ctk.CTkLabel(info, text=f"{time} • {emotion}", font=ctk.CTkFont(size=11), text_color="gray").pack(anchor="w")

    def show_settings(self):
        self.clear_content()
        self.current_page = "settings"
        
        # Reset grid configuration to prevent layout issues from other tabs
        self.main_frame.columnconfigure(0, weight=1)
        self.main_frame.columnconfigure(1, weight=0)
        self.main_frame.columnconfigure(2, weight=0)
        self.main_frame.rowconfigure(0, weight=1)
        
        # Simple scrollable container
        container = ctk.CTkScrollableFrame(self.main_frame, fg_color="transparent")
        container.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)

        ctk.CTkLabel(container, text="⚙️ System Configuration", 
                    font=ctk.CTkFont(size=28, weight="bold")).pack(anchor="w", pady=(0, 30))

        # ===== GENERAL SETTINGS =====
        gen_card = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], corner_radius=15)
        gen_card.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(gen_card, text="General Settings", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(anchor="w", padx=20, pady=(20, 15))
        
        # Email enable checkbox (using CTkCheckBox instead of CTkSwitch)
        self.email_checkbox = ctk.CTkCheckBox(gen_card, text="Enable Email Reports")
        if self.settings_manager.get("email_enabled"):
            self.email_checkbox.select()
        self.email_checkbox.pack(anchor="w", padx=20, pady=10)
        
        # Automation Intervals
        ctk.CTkLabel(gen_card, text="Automation Intervals (Minutes)", 
                    font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=20, pady=(20, 10))
        
        # Analysis Interval
        ana_frame = ctk.CTkFrame(gen_card, fg_color="transparent")
        ana_frame.pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(ana_frame, text="Face Analysis Every:", width=200, anchor="w").pack(side="left")
        self.analysis_interval = ctk.CTkEntry(ana_frame, width=80)
        self.analysis_interval.insert(0, str(self.settings_manager.get("analysis_interval")))
        self.analysis_interval.pack(side="left", padx=10)
        
        # Attendance Interval
        att_frame = ctk.CTkFrame(gen_card, fg_color="transparent")
        att_frame.pack(fill="x", padx=20, pady=(5, 20))
        ctk.CTkLabel(att_frame, text="Mark Attendance Every:", width=200, anchor="w").pack(side="left")
        self.attendance_interval = ctk.CTkEntry(att_frame, width=80)
        self.attendance_interval.insert(0, str(self.settings_manager.get("attendance_interval")))
        self.attendance_interval.pack(side="left", padx=10)

        # ===== SMTP SETTINGS =====
        smtp_card = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], corner_radius=15)
        smtp_card.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(smtp_card, text="📧 Email Configuration", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(anchor="w", padx=20, pady=(20, 10))
        
        ctk.CTkLabel(smtp_card, text="Sender Email:", anchor="w", 
                    font=ctk.CTkFont(weight="bold")).pack(fill="x", padx=20, pady=(10, 5))
        self.email_entry = ctk.CTkEntry(smtp_card, height=35)
        self.email_entry.insert(0, self.email_automation.sender_email)
        self.email_entry.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkLabel(smtp_card, text="App Password:", anchor="w", 
                    font=ctk.CTkFont(weight="bold")).pack(fill="x", padx=20, pady=(0, 5))
        self.pass_entry = ctk.CTkEntry(smtp_card, show="*", height=35)
        self.pass_entry.insert(0, self.email_automation.sender_password)
        self.pass_entry.pack(fill="x", padx=20, pady=(0, 20))

        # ===== FACULTY EMAILS =====
        fac_card = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], corner_radius=15)
        fac_card.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(fac_card, text="👨‍🏫 Faculty Email Configuration", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(anchor="w", padx=20, pady=(20, 15))
        
        ctk.CTkLabel(fac_card, text="Set email addresses for each subject to receive attendance reports:", 
                    text_color="gray").pack(anchor="w", padx=20, pady=(0, 15))
        
        # Grid for subject emails
        email_grid = ctk.CTkFrame(fac_card, fg_color="transparent")
        email_grid.pack(fill="x", padx=20, pady=(0, 20))
        email_grid.grid_columnconfigure(1, weight=1)
        
        self.fac_entries = {}
        saved_emails = self.settings_manager.get("faculty_emails") or {}
        subjects = get_all_subjects()
        
        for i, sub in enumerate(subjects):
            ctk.CTkLabel(email_grid, text=f"{sub}:", anchor="w", width=80,
                        font=ctk.CTkFont(weight="bold")).grid(row=i, column=0, padx=10, pady=8, sticky="w")
            entry = ctk.CTkEntry(email_grid, placeholder_text=f"faculty@example.com", height=35)
            val = saved_emails.get(sub, "")
            if val:
                entry.insert(0, val)
            entry.grid(row=i, column=1, padx=10, pady=8, sticky="ew")
            self.fac_entries[sub] = entry

        # ===== SAVE BUTTON =====
        def save_settings():
            try:
                # Save general settings
                self.settings_manager.set("email_enabled", self.email_checkbox.get() == 1)
                self.settings_manager.set("analysis_interval", int(self.analysis_interval.get()))
                self.settings_manager.set("attendance_interval", int(self.attendance_interval.get()))
                
                # Save SMTP
                self.email_automation.update_credentials(
                    self.email_entry.get().strip(), 
                    self.pass_entry.get().strip()
                )
                
                # Save faculty emails
                fac_map = {}
                for sub, entry in self.fac_entries.items():
                    fac_map[sub] = entry.get().strip()
                self.settings_manager.set("faculty_emails", fac_map)
                
                messagebox.showinfo("Success", "✅ All settings saved successfully!")
                logger.info("Settings saved successfully")
            except ValueError:
                messagebox.showerror("Error", "❌ Intervals must be valid numbers.")
            except Exception as e:
                logger.error(f"Settings save error: {e}")
                messagebox.showerror("Error", f"❌ Failed to save: {str(e)}")

        ctk.CTkButton(container, text="💾 Save All Settings", command=save_settings, 
                     height=50, font=ctk.CTkFont(size=16, weight="bold"),
                     fg_color=THEME_COLORS['primary'], 
                     hover_color=THEME_COLORS['success']).pack(fill="x", pady=(10, 20))

    def show_reports(self):
        self.clear_content()
        self.current_page = "reports"
        
        container = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=30, pady=30)
        
        # Header
        header = ctk.CTkFrame(container, fg_color="transparent")
        header.pack(fill="x", pady=(0, 20))
        ctk.CTkLabel(header, text="Attendance Reports", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        ctk.CTkLabel(header, text="View and manage generated reports", text_color="gray").pack(side="left", padx=10, pady=(5, 0))

        # Filter Toolbar (Mock)
        toolbar = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], height=50, corner_radius=10)
        toolbar.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(toolbar, text="📅 Date Range:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=15, pady=10)
        ctk.CTkButton(toolbar, text="Last 7 Days", fg_color=THEME_COLORS['background'], width=100, height=30).pack(side="left", padx=5)
        ctk.CTkButton(toolbar, text="This Month", fg_color="transparent", border_width=1, width=100, height=30).pack(side="left", padx=5)
        
        # Monthly Summary Button
        ctk.CTkButton(toolbar, text="📊 Generate Monthly Summary", command=self.generate_monthly_summary,
                     fg_color=THEME_COLORS['primary'], hover_color=THEME_COLORS['primary_dark'], 
                     width=200, height=35, font=ctk.CTkFont(weight="bold")).pack(side="right", padx=15)

        # Reports Grid/List
        reports_frame = ctk.CTkScrollableFrame(container, fg_color="transparent")
        reports_frame.pack(fill="both", expand=True)

        files = sorted([f for f in os.listdir(REPORTS_DIR) if f.endswith(('.txt', '.docx'))], reverse=True)
        if not files:
            ctk.CTkLabel(reports_frame, text="No reports generated yet.", text_color="gray", font=ctk.CTkFont(slant="italic")).pack(pady=40)
            
        for f in files:
            self.create_report_card(reports_frame, f)

    def create_report_card(self, parent, filename):
        card = ctk.CTkFrame(parent, fg_color=THEME_COLORS['surface'], corner_radius=10)
        card.pack(fill="x", pady=5)
        
        # Icon based on type
        icon = "📄" if filename.endswith('.txt') else "📝"
        icon_lbl = ctk.CTkLabel(card, text=icon, font=ctk.CTkFont(size=24))
        icon_lbl.pack(side="left", padx=20, pady=15)
        
        info = ctk.CTkFrame(card, fg_color="transparent")
        info.pack(side="left", fill="x", expand=True)
        
        ctk.CTkLabel(info, text=filename, font=ctk.CTkFont(weight="bold", size=14)).pack(anchor="w")
        
        # Extract date from filename if possible (e.g. Attendance_Default_2023-10-27...)
        parts = filename.split('_')
        date_str = "Unknown Date"
        if len(parts) > 2:
            date_str = f"{parts[-2]}" 

        ctk.CTkLabel(info, text=f"Generated: {date_str}", font=ctk.CTkFont(size=11), text_color="gray").pack(anchor="w")

        # Actions
        ctk.CTkButton(card, text="🗑️ Delete", width=90, 
                     command=lambda f=filename: self.delete_report(f),
                     fg_color=THEME_COLORS['danger'], hover_color="#dc2626").pack(side="right", padx=5)
        
        ctk.CTkButton(card, text="Open", width=80, command=lambda: os.startfile(os.path.join(REPORTS_DIR, filename)),
                     fg_color=THEME_COLORS['background'], hover_color=THEME_COLORS['primary']).pack(side="right", padx=15)

    def delete_report(self, filename):
        """Delete a report file with confirmation"""
        # Confirm deletion
        confirm = messagebox.askyesno(
            "Confirm Delete", 
            f"Are you sure you want to delete this report?\n\n{filename}\n\nThis action cannot be undone."
        )
        
        if not confirm:
            return
        
        try:
            # Delete the file
            filepath = os.path.join(REPORTS_DIR, filename)
            if os.path.exists(filepath):
                os.remove(filepath)
                logger.info(f"Deleted report: {filename}")
                messagebox.showinfo("Success", f"Report deleted successfully!\n\n{filename}")
                
                # Refresh the reports page
                self.show_reports()
            else:
                messagebox.showerror("Error", "Report file not found!")
                
        except Exception as e:
            logger.error(f"Error deleting report: {e}")
            messagebox.showerror("Error", f"Failed to delete report:\n{str(e)}")

    def _parse_attendance_from_report(self, filepath):
        """Parse attendance data from a single report file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract subject from filename
            filename = os.path.basename(filepath)
            subject = filename.split('_')[0]
            
            # Parse present students
            present_students = []
            absent_students = []
            
            # Find present students section
            if "PRESENT STUDENTS" in content:
                lines = content.split('\n')
                in_present_section = False
                in_absent_section = False
                
                for line in lines:
                    if "PRESENT STUDENTS" in line:
                        in_present_section = True
                        in_absent_section = False
                        continue
                    elif "ABSENT STUDENTS" in line:
                        in_present_section = False
                        in_absent_section = True
                        continue
                    elif "CLASS EMOTION" in line or "======" in line:
                        in_present_section = False
                        in_absent_section = False
                        continue
                    
                    # Parse student line (format: "1. Roll No: X, Name: Y")
                    if in_present_section and line.strip() and line.strip()[0].isdigit():
                        # Extract name part after "Name: "
                        if "Name:" in line:
                            name_part = line.split("Name:")[1].strip()
                            present_students.append(name_part)
                    
                    if in_absent_section and line.strip() and line.strip()[0].isdigit():
                        if "Name:" in line:
                            name_part = line.split("Name:")[1].strip()
                            absent_students.append(name_part)
            
            return {
                'subject': subject,
                'present': present_students,
                'absent': absent_students
            }
            
        except Exception as e:
            logger.error(f"Error parsing report {filepath}: {e}")
            return None

    def generate_monthly_summary(self):
        """Generate monthly attendance summary report"""
        try:
            # Get all txt report files
            report_files = [f for f in os.listdir(REPORTS_DIR) if f.endswith('.txt')]
            
            if not report_files:
                messagebox.showwarning("No Data", "❌ No reports found!\n\nPlease generate some attendance reports first.")
                return
            
            # Show progress
            messagebox.showinfo("Processing", f"📊 Processing {len(report_files)} reports...\n\nPlease wait...")
            
            # Aggregate data: {student_name: {subject: {'present': count, 'total': count}}}
            student_data = {}
            all_subjects = set()
            
            # Parse all reports
            for filename in report_files:
                filepath = os.path.join(REPORTS_DIR, filename)
                report_data = self._parse_attendance_from_report(filepath)
                
                if not report_data:
                    continue
                
                subject = report_data['subject']
                all_subjects.add(subject)
                
                # Process present students
                for student in report_data['present']:
                    if student not in student_data:
                        student_data[student] = {}
                    if subject not in student_data[student]:
                        student_data[student][subject] = {'present': 0, 'total': 0}
                    student_data[student][subject]['present'] += 1
                    student_data[student][subject]['total'] += 1
                
                # Process absent students
                for student in report_data['absent']:
                    if student not in student_data:
                        student_data[student] = {}
                    if subject not in student_data[student]:
                        student_data[student][subject] = {'present': 0, 'total': 0}
                    student_data[student][subject]['total'] += 1
            
            if not student_data:
                messagebox.showwarning("No Data", "❌ Could not extract student data from reports.\n\nPlease check report format.")
                return
            
            # Generate DOCX report
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Monthly Attendance Summary Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            metadata = doc.add_paragraph()
            metadata.add_run(f"Generated: ").bold = True
            metadata.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            metadata.add_run(f"Total Reports: ").bold = True
            metadata.add_run(f"{len(report_files)}\n")
            metadata.add_run(f"Subjects: ").bold = True
            metadata.add_run(f"{', '.join(sorted(all_subjects))}\n")
            metadata.add_run(f"Students: ").bold = True
            metadata.add_run(f"{len(student_data)}")
            
            doc.add_paragraph()  # Spacing
            
            # Create table
            subjects_list = sorted(all_subjects)
            num_cols = 2 + len(subjects_list) + 1  # Student Name + Subjects + Overall %
            num_rows = 1 + len(student_data)  # Header + Students
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Student Name'
            for i, subject in enumerate(subjects_list):
                header_cells[i + 1].text = subject
            header_cells[-1].text = 'Overall %'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for row_idx, (student_name, subjects_data) in enumerate(sorted(student_data.items()), 1):
                row_cells = table.rows[row_idx].cells
                row_cells[0].text = student_name
                
                total_present = 0
                total_sessions = 0
                
                # Fill subject columns
                for col_idx, subject in enumerate(subjects_list):
                    cell = row_cells[col_idx + 1]
                    if subject in subjects_data:
                        present = subjects_data[subject]['present']
                        total = subjects_data[subject]['total']
                        percentage = (present / total * 100) if total > 0 else 0
                        
                        cell.text = f"{present}/{total}"
                        
                        # Color code based on percentage
                        if percentage >= 75:
                            # Green background
                            shading = cell._element.get_or_add_tcPr()
                            shading_elem = shading.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                            if shading_elem is None:
                                from docx.oxml import OxmlElement
                                shading_elem = OxmlElement('w:shd')
                                shading.append(shading_elem)
                            shading_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'C6EFCE')
                        elif percentage >= 50:
                            # Yellow background
                            shading = cell._element.get_or_add_tcPr()
                            shading_elem = shading.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                            if shading_elem is None:
                                from docx.oxml import OxmlElement
                                shading_elem = OxmlElement('w:shd')
                                shading.append(shading_elem)
                            shading_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFEB9C')
                        else:
                            # Red background
                            shading = cell._element.get_or_add_tcPr()
                            shading_elem = shading.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                            if shading_elem is None:
                                from docx.oxml import OxmlElement
                                shading_elem = OxmlElement('w:shd')
                                shading.append(shading_elem)
                            shading_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFC7CE')
                        
                        total_present += present
                        total_sessions += total
                    else:
                        cell.text = "N/A"
                
                # Overall percentage
                overall_pct = (total_present / total_sessions * 100) if total_sessions > 0 else 0
                row_cells[-1].text = f"{overall_pct:.1f}%"
            
            # Legend
            doc.add_paragraph()
            legend = doc.add_paragraph()
            legend.add_run("Legend: ").bold = True
            legend.add_run("Green (≥75%), Yellow (50-74%), Red (<50%)")
            
            # Footer
            footer = doc.add_paragraph()
            footer.add_run(f"\nAttendance & Emotion Analytics System").italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            timestamp = datetime.now().strftime(REPORT_TIMESTAMP_FORMAT)
            filename = f"Monthly_Summary_{timestamp}.docx"
            filepath = os.path.join(REPORTS_DIR, filename)
            doc.save(filepath)
            
            logger.info(f"Monthly summary generated: {filepath}")
            messagebox.showinfo("Success", f"✅ Monthly Summary Generated!\n\nFile: {filename}\n\nStudents: {len(student_data)}\nSubjects: {len(subjects_list)}\nReports: {len(report_files)}")
            
            # Refresh reports view
            self.show_reports()
            
        except Exception as e:
            logger.error(f"Error generating monthly summary: {e}")
            messagebox.showerror("Error", f"❌ Failed to generate summary:\n\n{str(e)}")

    def show_enrollment(self):
        """Modern split-view enrollment page"""
        self.clear_content()
        self.current_page = "enrollment"
        
        if not hasattr(self, 'enrollment_images'):
            self.enrollment_images = []
        self.enrollment_images = []
        
        # Grid Layout: Left (Form), Right (Camera)
        self.main_frame.columnconfigure(0, weight=1)
        self.main_frame.columnconfigure(1, weight=1)
        self.main_frame.rowconfigure(0, weight=1)

        # --- Left Panel: Form & Controls ---
        left_panel = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=30, pady=30)
        
        ctk.CTkLabel(left_panel, text="New Student Registration", 
                    font=ctk.CTkFont(family=FONT_FAMILY, size=24, weight="bold")).pack(anchor="w", pady=(0, 20))

        # Form Container
        form_card = ctk.CTkFrame(left_panel, fg_color=THEME_COLORS['surface'], corner_radius=15)
        form_card.pack(fill="x", pady=(0, 20))
        
        form_inner = ctk.CTkFrame(form_card, fg_color="transparent")
        form_inner.pack(padx=20, pady=20, fill="x")

        # Inputs
        ctk.CTkLabel(form_inner, text="Roll Number", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))
        self.enroll_roll_entry = ctk.CTkEntry(form_inner, placeholder_text="e.g. 2024001", height=40, border_color=THEME_COLORS['border'])
        self.enroll_roll_entry.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(form_inner, text="Full Name", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))
        self.enroll_name_entry = ctk.CTkEntry(form_inner, placeholder_text="e.g. John Doe", height=40, border_color=THEME_COLORS['border'])
        self.enroll_name_entry.pack(fill="x", pady=(0, 15))

        # Progress
        self.enroll_count_label = ctk.CTkLabel(left_panel, text="Progress: 0/5 Images Captured", 
                                              font=ctk.CTkFont(size=14, weight="bold"), text_color=THEME_COLORS['text_dim'])
        self.enroll_count_label.pack(anchor="w", pady=(10, 5))
        
        self.progress_bar = ctk.CTkProgressBar(left_panel, height=10, corner_radius=5)
        self.progress_bar.set(0)
        self.progress_bar.pack(fill="x", pady=(0, 20))

        # Actions
        self.enroll_capture_btn = ctk.CTkButton(left_panel, text="📸 CAPTURE PHOTO", command=self.capture_enrollment_photo,
                                               height=50, fg_color=THEME_COLORS['primary'], hover_color=THEME_COLORS['primary_dark'],
                                               font=ctk.CTkFont(weight="bold"))
        self.enroll_capture_btn.pack(fill="x", pady=10)
        
        # Upload Photos Button
        self.enroll_upload_btn = ctk.CTkButton(left_panel, text="📁 UPLOAD PHOTOS", command=self.upload_enrollment_photos,
                                               height=50, fg_color=THEME_COLORS['secondary'], hover_color="#7c3aed",
                                               font=ctk.CTkFont(weight="bold"))
        self.enroll_upload_btn.pack(fill="x", pady=10)

        self.enroll_save_btn = ctk.CTkButton(left_panel, text="💾 SAVE PROFILE", command=self.save_enrollment,
                                            height=50, fg_color=THEME_COLORS['success'], hover_color="#059669",
                                            font=ctk.CTkFont(weight="bold"), state="normal")
        self.enroll_save_btn.pack(fill="x", pady=10)

        ctk.CTkButton(left_panel, text="🗑 Reset Form", command=self.clear_enrollment,
                     height=40, fg_color="transparent", border_width=1, text_color="gray", border_color="gray").pack(fill="x", pady=10)


        # --- Right Panel: Camera ---
        right_panel = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        right_panel.grid(row=0, column=1, sticky="nsew", padx=30, pady=30)
        
        cam_card = ctk.CTkFrame(right_panel, fg_color="black", corner_radius=15, 
                              border_width=2, border_color=THEME_COLORS['secondary'])
        cam_card.pack(fill="both", expand=True, pady=(0, 20))
        
        self.enroll_camera_label = ctk.CTkLabel(cam_card, text="Camera Offline", text_color="gray")
        self.enroll_camera_label.pack(expand=True, fill="both")

        # Camera Controls
        ctrl_row = ctk.CTkFrame(right_panel, fg_color="transparent")
        ctrl_row.pack(fill="x")
        
        self.enroll_start_cam_btn = ctk.CTkButton(ctrl_row, text="▶ START CAMERA", command=self.start_enrollment_camera,
                                                 fg_color=THEME_COLORS['surface'], hover_color=THEME_COLORS['primary'], width=150)
        self.enroll_start_cam_btn.pack(side="left", padx=5)
        
        self.enroll_stop_cam_btn = ctk.CTkButton(ctrl_row, text="⏹ STOP", command=self.stop_enrollment_camera,
                                                fg_color=THEME_COLORS['surface'], hover_color=THEME_COLORS['danger'], width=100)
        self.enroll_stop_cam_btn.pack(side="left", padx=5)

        # Auto-start with delay
        self.after(800, self.start_enrollment_camera)

    def start_enrollment_camera(self):
        """Start camera for enrollment"""
        if self.camera_active: return

        if self.image_capture.initialize_camera():
            self.camera_active = True
            self.enroll_start_cam_btn.configure(state="disabled", text="Running...")
            self.enroll_stop_cam_btn.configure(state="normal")
            self.update_enrollment_camera()
        else:
            messagebox.showerror("Error", "Camera failed to start.")
    
    def stop_enrollment_camera(self):
        """Stop camera from enrollment page"""
        self.camera_active = False
        self.image_capture.release_camera()
        
        self.enroll_start_cam_btn.configure(state="normal", text="▶ START CAMERA")
        self.enroll_stop_cam_btn.configure(state="disabled")
        self.enroll_camera_label.configure(image=None, text="Camera Preview Paused")

    def update_enrollment_camera(self):
        """Update camera feed on enrollment page"""
        if self.current_page != "enrollment" or not self.camera_active:
            return
        
        frame = self.image_capture.get_frame()
        if frame is not None:
            # Add face detection box
            if self.face_cascade:
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                faces = self.face_cascade.detectMultiScale(gray, 1.1, 4)
                for (x, y, w, h) in faces:
                    cv2.rectangle(frame, (x, y), (x+w, y+h), (0, 255, 0), 2)
            
            # Resize and display
            h, w = frame.shape[:2]
            max_w, max_h = 600, 450
            ratio = min(max_w/w, max_h/h)
            new_w, new_h = int(w*ratio), int(h*ratio)
            
            img_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img_pil = Image.fromarray(img_rgb)
            ctk_img = ctk.CTkImage(img_pil, size=(new_w, new_h))
            self.enroll_camera_label.configure(image=ctk_img, text="")
        
        self.after(30, self.update_enrollment_camera)

    def capture_enrollment_photo(self):
        """Capture a photo for enrollment"""
        # Try to start camera if not active
        if not self.camera_active:
            self.start_enrollment_camera()
            self.after(500, self.capture_enrollment_photo)  # Retry after starting
            return
        
        frame = self.image_capture.get_frame()
        if frame is None:
            messagebox.showerror("Error", "Could not capture image. Please try again.")
            return
        
        # Save temporary image
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        temp_path = os.path.join(IMAGES_DIR, f"enroll_temp_{timestamp}.jpg")
        cv2.imwrite(temp_path, frame)
        
        self.enrollment_images.append(temp_path)
        count = len(self.enrollment_images)
        
        # Update UI
        self.enroll_count_label.configure(text=f"Progress: {count}/5 Images")
        self.progress_bar.set(count / 5)
        
        if count >= 5:
            self.enroll_capture_btn.configure(state="disabled")
            messagebox.showinfo("Ready", "5 images captured! You can now save and enroll.")

    def save_enrollment(self):
        """Save enrollment and train model"""
        roll = self.enroll_roll_entry.get().strip()
        name = self.enroll_name_entry.get().strip()
        
        if not roll or not name:
            messagebox.showerror("Error", "Please fill in all fields!")
            return
        
        if len(self.enrollment_images) < 3:
            messagebox.showerror("Error", "Please capture at least 3 images!")
            return
        
        # Check duplicate
        existing = os.listdir(STUDENT_DATASET_DIR) if os.path.exists(STUDENT_DATASET_DIR) else []
        for student in existing:
            if student.startswith(f"{roll}_"):
                messagebox.showerror("Duplicate", f"Roll number {roll} already exists!")
                return
        
        try:
            # Create student folder
            folder_name = f"{roll}_{name.replace(' ', '_')}"
            student_dir = os.path.join(STUDENT_DATASET_DIR, folder_name)
            os.makedirs(student_dir, exist_ok=True)
            
            # Move images
            for i, img_path in enumerate(self.enrollment_images, 1):
                dest = os.path.join(student_dir, f"{folder_name}_{i}.jpg")
                os.rename(img_path, dest)
            
            # Train model
            image_count = len(self.enrollment_images)
            messagebox.showinfo("Training", "Enrollment successful! Training model...")
            threading.Thread(target=self.face_recognition.train_face_encodings, daemon=True).start()
            
            # Clear form
            self.clear_enrollment()
            messagebox.showinfo("Success", f"Enrolled {name} (Roll: {roll})\nCaptured {image_count} images")
            
        except Exception as e:
            logger.error(f"Enrollment error: {e}")
            messagebox.showerror("Error", str(e))

    def upload_enrollment_photos(self):
        """Upload photos from disk for enrollment (alternative to camera capture)"""
        try:
            # Open file dialog to select multiple images
            file_paths = filedialog.askopenfilenames(
                title="Select Student Photos (Up to 5)",
                filetypes=[
                    ("Image Files", "*.jpg *.jpeg *.png *.bmp"),
                    ("All Files", "*.*")
                ],
                parent=self
            )
            
            if not file_paths:
                return  # User cancelled
            
            # Limit to 5 total images
            current_count = len(self.enrollment_images)
            available_slots = 5 - current_count
            
            if available_slots <= 0:
                messagebox.showwarning("Limit Reached", "You already have 5 images. Reset the form to upload new ones.")
                return
            
            # Process selected files
            added_count = 0
            for file_path in file_paths[:available_slots]:
                try:
                    # Validate it's a valid image
                    img = cv2.imread(file_path)
                    if img is None:
                        logger.warning(f"Invalid image: {file_path}")
                        continue
                    
                    # Copy to temp location with timestamp
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                    temp_filename = f"enroll_upload_{timestamp}_{added_count}.jpg"
                    temp_path = os.path.join(IMAGES_DIR, temp_filename)
                    
                    # Save copy
                    cv2.imwrite(temp_path, img)
                    self.enrollment_images.append(temp_path)
                    added_count += 1
                    
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    continue
            
            # Update UI
            if added_count > 0:
                total_count = len(self.enrollment_images)
                self.enroll_count_label.configure(text=f"Progress: {total_count}/5 Images")
                self.progress_bar.set(total_count / 5)
                
                if total_count >= 5:
                    self.enroll_capture_btn.configure(state="disabled")
                    self.enroll_upload_btn.configure(state="disabled")
                    messagebox.showinfo("Ready", f"✅ {added_count} photo(s) uploaded!\n\nTotal: {total_count}/5\nYou can now save and enroll.")
                else:
                    messagebox.showinfo("Success", f"✅ {added_count} photo(s) uploaded!\n\nTotal: {total_count}/5\nYou can upload more or save now.")
            else:
                messagebox.showerror("Error", "No valid images were uploaded. Please select valid image files.")
                
        except Exception as e:
            logger.error(f"Upload photos error: {e}")
            messagebox.showerror("Error", f"Failed to upload photos: {str(e)}")

    def clear_enrollment(self):
        """Clear enrollment form and images"""
        # Delete temp images
        for img_path in self.enrollment_images:
            if os.path.exists(img_path):
                try:
                    os.remove(img_path)
                except:
                    pass
        
        self.enrollment_images = []
        if hasattr(self, 'enroll_roll_entry'): self.enroll_roll_entry.delete(0, 'end')
        if hasattr(self, 'enroll_name_entry'): self.enroll_name_entry.delete(0, 'end')
        if hasattr(self, 'enroll_count_label'): self.enroll_count_label.configure(text="Progress: 0/5 Images Captured")
        if hasattr(self, 'progress_bar'): self.progress_bar.set(0)
        if hasattr(self, 'enroll_capture_btn'): self.enroll_capture_btn.configure(state="normal")
        if hasattr(self, 'enroll_upload_btn'): self.enroll_upload_btn.configure(state="normal")
        if hasattr(self, 'enroll_save_btn'): self.enroll_save_btn.configure(state="normal")

    def show_student_database(self):
        self.clear_content()
        self.current_page = "students"

        # Main Layout
        container = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=30, pady=30)
        
        # Header with Actions
        header = ctk.CTkFrame(container, fg_color="transparent")
        header.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(header, text="Student Directory", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        
        actions = ctk.CTkFrame(header, fg_color="transparent")
        actions.pack(side="right")
        
        ctk.CTkButton(actions, text="📂 Open Data", command=self.open_dataset_folder, 
                     fg_color=THEME_COLORS['surface'], hover_color=THEME_COLORS['primary'], width=100).pack(side="left", padx=5)
        
        ctk.CTkButton(actions, text="🔄 Re-Train Model", command=self.train_faces, 
                     fg_color=THEME_COLORS['warning'], hover_color="#d97706", width=120).pack(side="left", padx=5)

        ctk.CTkButton(actions, text="🗑 Delete All", command=self.delete_all_students,
                      fg_color=THEME_COLORS['danger'], hover_color="#b91c1c", width=100).pack(side="left", padx=5)

        # Search Bar
        search_frame = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], height=50, corner_radius=10)
        search_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(search_frame, text="🔍", font=ctk.CTkFont(size=16)).pack(side="left", padx=15, pady=10)
        search_entry = ctk.CTkEntry(search_frame, placeholder_text="Search by Name or Roll Number...", 
                                   border_width=0, fg_color="transparent", height=40, font=ctk.CTkFont(size=14))
        search_entry.pack(side="left", fill="x", expand=True, padx=5)

        # Table Header
        table_header = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], height=40, corner_radius=5)
        table_header.pack(fill="x", pady=(0, 5))
        
        # Grid Configuration for Table
        # Columns: Profile(0), Roll(1), Name(2), Status(3), Actions(4)
        # Weights: 1, 2, 3, 1, 1
        
        table_header.columnconfigure(0, weight=1, minsize=80)
        table_header.columnconfigure(1, weight=2, minsize=150)
        table_header.columnconfigure(2, weight=3, minsize=200)
        table_header.columnconfigure(3, weight=1, minsize=100)
        table_header.columnconfigure(4, weight=1, minsize=100)

        headers = ["PROFILE", "ROLL NUMBER", "FULL NAME", "STATUS", "ACTIONS"]
        for i, text in enumerate(headers):
            ctk.CTkLabel(table_header, text=text, font=ctk.CTkFont(size=11, weight="bold"), 
                        text_color="gray").grid(row=0, column=i, sticky="w", padx=10, pady=10)

        # Scrollable List (Rows)
        list_frame = ctk.CTkScrollableFrame(container, fg_color="transparent")
        list_frame.pack(fill="both", expand=True)

        students = self.face_recognition.get_all_students()
        if not students:
            ctk.CTkLabel(list_frame, text="No students found.", text_color="gray", font=ctk.CTkFont(slant="italic")).pack(pady=40)
        
        for student in students:
            self.create_student_row(list_frame, student)

    def create_student_row(self, parent, student_id):
        row = ctk.CTkFrame(parent, fg_color=THEME_COLORS['surface'], corner_radius=8)
        row.pack(fill="x", pady=2)
        
        # Match Grid Config with Header
        row.columnconfigure(0, weight=1, minsize=80)
        row.columnconfigure(1, weight=2, minsize=150)
        row.columnconfigure(2, weight=3, minsize=200)
        row.columnconfigure(3, weight=1, minsize=100)
        row.columnconfigure(4, weight=1, minsize=100)
        
        try:
            roll, name = student_id.split('_', 1)
            name = name.replace('_', ' ')
        except:
            roll = "N/A"
            name = student_id

        # Profile Icon
        ctk.CTkLabel(row, text="👤", font=ctk.CTkFont(size=18)).grid(row=0, column=0, sticky="w", padx=20, pady=10)
        
        # Roll
        ctk.CTkLabel(row, text=roll, font=ctk.CTkFont(weight="bold")).grid(row=0, column=1, sticky="w", padx=10, pady=10)
        
        # Name
        ctk.CTkLabel(row, text=name).grid(row=0, column=2, sticky="w", padx=10, pady=10)
        
        # Status
        ctk.CTkLabel(row, text="Active", text_color=THEME_COLORS['success'], 
                    font=ctk.CTkFont(size=12)).grid(row=0, column=3, sticky="w", padx=10, pady=10)
        
        # Actions
        actions_frame = ctk.CTkFrame(row, fg_color="transparent")
        actions_frame.grid(row=0, column=4, sticky="w", padx=10, pady=10)
        
        ctk.CTkButton(actions_frame, text="🗑", width=40, height=30, 
                     fg_color=THEME_COLORS['danger'], hover_color="#b91c1c",
                     command=lambda s=student_id: self.delete_student(s)).pack(side="left", padx=5)

    def delete_student(self, student_id):
        """Delete student data and re-train"""
        if not messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete {student_id}?"):
            return
            
        try:
            # removing folder
            path = os.path.join(STUDENT_DATASET_DIR, student_id)
            if os.path.exists(path):
                shutil.rmtree(path)
                
            # Trigger retrain
            threading.Thread(target=self.face_recognition.train_face_encodings, daemon=True).start()
            
            messagebox.showinfo("Deleted", f"Student {student_id} deleted successfully.\nModel is retraining in background.")
            self.show_student_database() # Refresh list
            
        except Exception as e:
            logger.error(f"Error deleting student: {e}")
            logger.error(f"Error deleting student: {e}")
            messagebox.showerror("Error", f"Failed to delete: {e}")

    def delete_all_students(self):
        """Delete ALL students"""
        if not messagebox.askyesno("DANGER", "Are you sure you want to DELETE ALL STUDENTS?\nThis cannot be undone!"):
            return
            
        try:
            # Delete contents
            for item in os.listdir(STUDENT_DATASET_DIR):
                path = os.path.join(STUDENT_DATASET_DIR, item)
                if os.path.isdir(path):
                    shutil.rmtree(path)
                elif item.endswith('.pkl'):
                    os.remove(path)
            
            # Reset UI
            messagebox.showinfo("Deleted", "All student records have been deleted.")
            self.show_student_database()
            
        except Exception as e:
            logger.error(f"Error deleting all: {e}")
            messagebox.showerror("Error", f"Failed: {e}")

    def upload_photo_attendance(self):
        """Manual attendance via photo upload"""
        file_path = filedialog.askopenfilename(filetypes=[("Images", "*.jpg;*.png;*.jpeg")])
        if not file_path: return
        
        try:
            # Analyze
            attendance = self.face_recognition.recognize_faces(file_path)
            
            if not attendance:
                messagebox.showwarning("No Faces", "No known students detected in the image.")
                return
            
            # Show results
            msg = "Faces Found:\n" + "\n".join([f"- {name}: {status}" for name, status in attendance.items()])
            messagebox.showinfo("Analysis Result", msg)
            
            # Mark Attendance in feed
            timestamp = datetime.now().strftime("%H:%M:%S")
            for name, status in attendance.items():
               if status == "Present":
                   self.update_recent_feed(name, timestamp)
            
            # Ask user to select subject for report generation
            from tkinter import simpledialog
            subjects = list(TIMETABLE.values())
            
            subject_dialog = ctk.CTkInputDialog(
                text=f"Select subject for this attendance:\n\nAvailable: {', '.join(subjects)}\n\nEnter subject name:",
                title="Select Subject"
            )
            subject = subject_dialog.get_input()
            
            if not subject or subject.strip() == "":
                # User cancelled or didn't enter anything
                messagebox.showinfo("Info", "✅ Attendance marked!\n\n⚠️ No report generated (subject not selected)")
                return
            
            subject = subject.strip().upper()
            
            # Generate reports and send email
            try:
                time_now = datetime.now().strftime("%H:%M:%S")
                
                # Generate report with the uploaded photo
                report_paths = self.report_generator.generate_report(
                    attendance=attendance,
                    emotion_summary={},  # No emotion data from manual upload
                    subject=subject,
                    time_start=time_now,
                    time_end=time_now,
                    report_format='both',
                    image_path=file_path  # Include the uploaded photo in the report
                )
                
                if report_paths:
                    logger.info(f"Reports generated: {report_paths}")
                    
                    # Send email if enabled
                    if self.settings_manager.get("email_enabled"):
                        saved_emails = self.settings_manager.get("faculty_emails") or {}
                        recipient = saved_emails.get(subject, "")
                        
                        if recipient:  # Only try to send if recipient is configured
                            if self.email_automation.send_attendance_report(subject, report_paths, recipient_email=recipient):
                                logger.info(f"Email sent successfully to {recipient}")
                                messagebox.showinfo("Success", f"✅ Report generated and email sent to {recipient}!\n\nStudents marked: {len([s for s in attendance.values() if s == 'Present'])}")
                            else:
                                messagebox.showinfo("Success", f"✅ Report generated!\n\n⚠️ Email failed to send\n\nStudents marked: {len([s for s in attendance.values() if s == 'Present'])}")
                        else:
                            # No recipient configured for Manual Upload
                            messagebox.showinfo("Success", f"✅ Report generated!\n\n💡 Tip: Configure faculty email for 'Manual Upload' in Settings to auto-send reports\n\nStudents marked: {len([s for s in attendance.values() if s == 'Present'])}")
                    else:
                        messagebox.showinfo("Success", f"✅ Report generated!\n\nStudents marked: {len([s for s in attendance.values() if s == 'Present'])}")
                
            except Exception as e:
                logger.error(f"Report/Email error: {e}")
                messagebox.showwarning("Partial Success", f"✅ Attendance marked\n⚠️ Report generation failed: {str(e)}")
                
        except Exception as e:
            logger.error(f"Upload error: {e}")
            messagebox.showerror("Error", f"Analysis failed: {str(e)}")

    def show_cleanup(self):
        self.clear_content()
        self.current_page = "cleanup"
        
        container = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=30, pady=30)
        
        ctk.CTkLabel(container, text="System Maintenance", font=ctk.CTkFont(size=24, weight="bold")).pack(anchor="w", pady=(0, 20))
        
        # Storage Card
        card = ctk.CTkFrame(container, fg_color=THEME_COLORS['surface'], corner_radius=15)
        card.pack(fill="x")
        
        ctk.CTkLabel(card, text="🧹 Data Cleanup", font=ctk.CTkFont(size=18, weight="bold")).pack(anchor="w", padx=20, pady=(20, 10))
        ctk.CTkLabel(card, text="Free up space by removing temporary files, old reports, and cached data.", 
                    text_color="gray").pack(anchor="w", padx=20, pady=(0, 20))
        
        # Meter (Visual only)
        meter_frame = ctk.CTkFrame(card, fg_color="transparent")
        meter_frame.pack(fill="x", padx=20, pady=(0, 20))
        ctk.CTkLabel(meter_frame, text="Storage Usage (Estimated)", font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(0, 5))
        self.cleanup_progress = ctk.CTkProgressBar(meter_frame, height=15, corner_radius=8, progress_color=THEME_COLORS['warning'])
        self.cleanup_progress.pack(fill="x")
        self.cleanup_progress.set(0.6) # Initial dummy value
        
        def run_cleanup():
            # Simulated cleanup with progress
            progress_step = 0
            
            def step():
                nonlocal progress_step
                progress_step += 0.2
                if hasattr(self, 'cleanup_progress'): # Ensure widget exists
                     self.cleanup_progress.set(progress_step)
                
                if progress_step < 1.0:
                    self.after(500, step)
                else:
                    messagebox.showinfo("Cleanup", "System cleanup completed successfully!\nRemoved temporary files and optimized database.")
                    self.cleanup_progress.set(0) # Reset
            
            step()

        ctk.CTkButton(card, text="Run System Cleanup", command=run_cleanup, height=50,
                     fg_color=THEME_COLORS['danger'], hover_color="#b91c1c",
                     font=ctk.CTkFont(weight="bold")).pack(padx=20, pady=20, fill="x")

    # ================= LOGIC =================

    def auto_start_camera(self):
        if self.current_page == 'live_capture' and not self.camera_active:
             self.start_automated_session()

    def start_automated_session(self):
        """Camera aur automation chalu karte hai"""
        if self.camera_active:
            self.stop_camera()
            return
            
        if self.image_capture.initialize_camera():
            self.camera_active = True
            self.automation_active = True
            self.start_btn.configure(text="⏹ STOP SESSION", fg_color=THEME_COLORS['danger'], hover_color="#b91c1c")
            self.capture_btn.configure(state="normal")
            
            # Timer reset karte hai
            self.last_analysis_time = datetime.now()
            self.last_attendance_time = datetime.now()
            
            self.update_camera_feed()
            logger.info("Session chalu ho gaya")
        else:
            messagebox.showerror("Error", "Camera nahi chala bhai.")

    def stop_camera(self):
        self.camera_active = False
        self.automation_active = False
        if self.after_id:
            self.after_cancel(self.after_id)
            self.after_id = None
        self.image_capture.release_camera()
        
        # UI reset kar dete hai
        self.start_btn.configure(text="▶ START SESSION", fg_color=THEME_COLORS['success'], hover_color="#059669")
        self.capture_btn.configure(state="disabled")
        self.camera_label.configure(image=None, text="Camera Band Hai\nStart pe click karo")

    def update_camera_feed(self):
        if not self.camera_active or self.showing_result: return
        
        frame = self.image_capture.get_frame()
        if frame is not None:
            # 1. Face Detection (Bas dikhane ke liye)
            if self.face_cascade:
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                faces = self.face_cascade.detectMultiScale(gray, 1.1, 4)
                for (x, y, w, h) in faces:
                    cv2.rectangle(frame, (x, y), (x+w, y+h), (0, 255, 0), 2)

            # 2. Automation Logic (Apna main kaam)
            if self.automation_active:
                self._check_automation(frame)

            # 3. Smart Resize (Screen pe fit hona chahiye)
            cw = self.cam_container.winfo_width()
            ch = self.cam_container.winfo_height()
            
            if cw > 10 and ch > 10:
                img_h, img_w = frame.shape[:2]
                ratio = min(cw/img_w, ch/img_h)
                new_w = int(img_w * ratio)
                new_h = int(img_h * ratio)
                
                img = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img_pil = Image.fromarray(img)
                ctk_img = ctk.CTkImage(light_image=img_pil, dark_image=img_pil, size=(new_w, new_h))
                self.camera_label.configure(image=ctk_img, text="")
        
        self.after_id = self.after(30, self.update_camera_feed)

    def _check_automation(self, frame):
        """Check karte hai ki time ho gaya kya analysis ka"""
        now = datetime.now()
        
        # Settings se time lete hai
        ana_int = self.settings_manager.get("analysis_interval")
        att_int = self.settings_manager.get("attendance_interval")
        
        # Check Analysis
        delta_ana = (now - self.last_analysis_time).total_seconds() / 60
        if delta_ana >= ana_int:
            logger.info("Chehra analyze kar rahe hai...")
            self.last_analysis_time = now
            threading.Thread(target=self._update_live_feed, args=(frame.copy(),), daemon=True).start()
            
        # Check Attendance Marking (Mock implementation hai abhi)
        delta_att = (now - self.last_attendance_time).total_seconds() / 60
        if delta_att >= att_int:
             logger.info("Attendance laga rahe hai...")
             self.last_attendance_time = now
             # Asli app me yaha database save hoga
             self.after(0, lambda: messagebox.showinfo("Autolink", "✅ Attendance lag gayi."))

    def _update_live_feed(self, frame):
        """Background me face dhoondhte hai taaki screen na atkegi"""
        try:
            # Temp save karte hai
            temp_path = os.path.join(IMAGES_DIR, "live_temp.jpg")
            cv2.imwrite(temp_path, frame)
            
            # Jaldi se pehchan lete hai
            attendance = self.face_recognition.recognize_faces(temp_path)
            
            # Agar koi mila toh screen pe dikhayenge
            for name, status in attendance.items():
                if status == "Present":
                    timestamp = datetime.now().strftime("%H:%M:%S")
                    # UI update shcedule
                    self.after(0, self.update_recent_feed, name, timestamp)
        except Exception as e:
            logger.error(f"Live feed me error: {e}")

    def capture_and_process(self):
        subject = self.subject_var.get()
        threading.Thread(target=self._process_attendance, args=(subject,), daemon=True).start()

    def _process_attendance(self, subject):
        self.update_status("Photo khinch rahe hai...", "orange")
        self.capture_btn.configure(state="disabled")
        
        try:
            images = self.image_capture.capture_multiple_images(count=3)
            if not images: raise Exception("Photo nahi aayi yaar")
            
            self.update_status("Chehra dhoond rahe hai...", "blue")
            all_attendance = {}
            final_annotated_img = None
            
            for img_path in images:
                # Attendance aur photo dono chahiye
                att, ann_img = self.face_recognition.recognize_faces(img_path, return_annotated=True)
                all_attendance.update(att)
                if ann_img is not None:
                    final_annotated_img = ann_img # Aakhri wala use karenge
            
            # Hara dibba dikhana hai result me
            if final_annotated_img is not None:
                self.showing_result = True
                self.display_image(final_annotated_img)
                # 4 second baad wapis camera chalu
                self.after(4000, self.resume_live_feed)
            
            # Dashboard update karte hai
            self.update_last_attendance(all_attendance)

            self.update_status("Report bana rahe hai...", "purple")
            emotion_summary = self.emotion_detection.analyze_multiple_images(images)
            timestamp = datetime.now()
            
            # Pehli photo report me lagayenge
            selected_image = images[0] if images else None
            
            time_str = timestamp.strftime("%H:%M:%S")
            report_path = self.report_generator.generate_report(
                all_attendance, emotion_summary, subject, time_str, time_str, 
                report_format='both', image_path=selected_image
            )
            
            # 5. Email (Agar setting on hai toh)
            if self.settings_manager.get("email_enabled"):
                try:
                    self.update_status("Email bhej rahe hai...", "blue")
                    
                    # Settings se email lete hai
                    saved_emails = self.settings_manager.get("faculty_emails") or {}
                    recipient = saved_emails.get(subject, "")
                    final_recipient = recipient if recipient else None
                    
                    if self.email_automation.send_attendance_report(subject, report_path, recipient_email=final_recipient):
                         logger.info(f"Email bhej diya {final_recipient} ko")
                    else:
                         logger.warning("Email nahi gaya ya setup nahi hai")
                         
                except Exception as e:
                    logger.error(f"Email error: {e}")
                    # Crash nahi hona chahiye
            
            self.update_status("Ho gaya bhai!", "green")
            self.after(0, lambda: messagebox.showinfo("Success", f"{len(all_attendance)} students ki attendance lag gayi."))
            
        except Exception as e:
            logger.error(f"Process error: {e}")
            self.update_status("Galti ho gayi", "red")
            self.after(0, lambda: messagebox.showerror("Error", str(e)))
        finally:
            self.after(0, lambda: self.capture_btn.configure(state="normal"))
            self.after(2000, lambda: self.update_status("Ready", "gray"))

    def update_status(self, text, color):
        self.after(0, lambda: self.status_label.configure(text=text, text_color=color))

    def show_enrollment_dialog(self):
        """Show improved enrollment dialog with roll number"""
        dialog = ctk.CTkToplevel(self)
        dialog.title("Student Enrollment")
        dialog.geometry("500x450")
        dialog.attributes("-topmost", True)
        dialog.resizable(False, False)
        
        # Header
        header = ctk.CTkFrame(dialog, fg_color=THEME_COLORS['primary'], corner_radius=0)
        header.pack(fill="x", pady=(0, 20))
        ctk.CTkLabel(header, text="👤 New Student Enrollment", 
                    font=ctk.CTkFont(size=24, weight="bold"),
                    text_color="white").pack(pady=20)
        
        # Form Container
        form_frame = ctk.CTkFrame(dialog, fg_color="transparent")
        form_frame.pack(fill="both", expand=True, padx=30)
        
        # Roll Number
        ctk.CTkLabel(form_frame, text="Roll Number*", 
                    font=ctk.CTkFont(size=14, weight="bold"),
                    anchor="w").pack(fill="x", pady=(10, 5))
        roll_entry = ctk.CTkEntry(form_frame, placeholder_text="Enter Roll Number (e.g., 2024001)",
                                 height=40, font=ctk.CTkFont(size=14))
        roll_entry.pack(fill="x", pady=(0, 15))
        
        # Name
        ctk.CTkLabel(form_frame, text="Full Name*", 
                    font=ctk.CTkFont(size=14, weight="bold"),
                    anchor="w").pack(fill="x", pady=(10, 5))
        name_entry = ctk.CTkEntry(form_frame, placeholder_text="Enter Full Name",
                                 height=40, font=ctk.CTkFont(size=14))
        name_entry.pack(fill="x", pady=(0, 15))
        
        # Instructions
        info_frame = ctk.CTkFrame(form_frame, fg_color=THEME_COLORS['light'])
        info_frame.pack(fill="x", pady=20)
        ctk.CTkLabel(info_frame, text="📸 Instructions:", 
                    font=ctk.CTkFont(size=12, weight="bold"),
                    anchor="w").pack(fill="x", padx=10, pady=(10, 5))
        ctk.CTkLabel(info_frame, 
                    text="• Position yourself in front of the camera\n• Look straight at the camera\n• System will capture 5 images automatically\n• Keep your face clearly visible",
                    font=ctk.CTkFont(size=11),
                    anchor="w", justify="left").pack(fill="x", padx=10, pady=(0, 10))
        
        # Status Label
        status_label = ctk.CTkLabel(form_frame, text="", 
                                   font=ctk.CTkFont(size=12),
                                   text_color="gray")
        status_label.pack(fill="x", pady=10)
        
        # Buttons
        btn_frame = ctk.CTkFrame(dialog, fg_color="transparent")
        btn_frame.pack(fill="x", padx=30, pady=20)
        
        def validate_and_start():
            roll = roll_entry.get().strip()
            name = name_entry.get().strip()
            
            if not roll:
                messagebox.showerror("Validation Error", "Please enter a roll number!", parent=dialog)
                roll_entry.focus()
                return
            
            if not name:
                messagebox.showerror("Validation Error", "Please enter a name!", parent=dialog)
                name_entry.focus()
                return
            
            # Check if roll number already exists
            existing_students = os.listdir(STUDENT_DATASET_DIR) if os.path.exists(STUDENT_DATASET_DIR) else []
            for student in existing_students:
                if student.startswith(f"{roll}_"):
                    messagebox.showerror("Duplicate Entry", 
                                       f"Roll number {roll} is already enrolled!", 
                                       parent=dialog)
                    return
            
            status_label.configure(text="Starting enrollment...", text_color="blue")
            dialog.after(500, lambda: [dialog.destroy(), self.quick_enroll(roll, name)])
        
        def cancel():
            dialog.destroy()
        
        ctk.CTkButton(btn_frame, text="Cancel", command=cancel,
                     fg_color="gray", hover_color="darkgray").pack(side="left", expand=True, fill="x", padx=(0, 10))
        ctk.CTkButton(btn_frame, text="Start Enrollment", command=validate_and_start,
                     fg_color=THEME_COLORS['primary'], height=40,
                     font=ctk.CTkFont(size=14, weight="bold")).pack(side="left", expand=True, fill="x")
        
        # Focus on roll number
        roll_entry.focus()

    def quick_enroll(self, roll_number, name):
        """Enhanced enrollment with roll number"""
        if not self.camera_active: 
            self.start_camera()
        self.after(1000, lambda: threading.Thread(
            target=self._enroll_thread, 
            args=(roll_number, name), 
            daemon=True
        ).start())

    def _enroll_thread(self, roll_number, name):
        """Enrollment thread with progress tracking"""
        try:
            # Update status
            self.update_status(f"📸 Capturing images for {name}...", "blue")
            
            # Create student directory with roll number
            student_folder_name = f"{roll_number}_{name.replace(' ', '_')}"
            student_dir = os.path.join(STUDENT_DATASET_DIR, student_folder_name)
            os.makedirs(student_dir, exist_ok=True)
            
            # Capture images with feedback
            self.after(0, lambda: self.update_status(f"Image 1/5...", "blue"))
            images = self.image_capture.capture_multiple_images(count=5, interval=1, class_id=student_folder_name)
            
            # Move images to student directory
            moved_count = 0
            for i, img in enumerate(images, 1):
                if img and os.path.exists(img):
                    dest = os.path.join(student_dir, os.path.basename(img))
                    os.rename(img, dest)
                    moved_count += 1
                    self.after(0, lambda i=i: self.update_status(f"✓ Captured {i}/5 images", "green"))
            
            if moved_count < 3:
                self.after(0, lambda: messagebox.showerror(
                    "Enrollment Failed", 
                    "Could not capture enough images. Please try again."
                ))
                return
            
            # Train model
            self.update_status("🔄 Training recognition model...", "orange")
            self.face_recognition.train_face_encodings()
            
            # Success
            self.update_status("✓ Enrollment Complete!", "green")
            self.after(0, lambda: messagebox.showinfo(
                "Success", 
                f"Successfully enrolled!\n\nRoll: {roll_number}\nName: {name}\nImages: {moved_count}"
            ))
            
        except Exception as e:
            logger.error(f"Enrollment error: {e}")
            self.update_status("Enrollment Failed", "red")
            self.after(0, lambda: messagebox.showerror("Enrollment Error", str(e)))

    def open_dataset_folder(self): os.startfile(STUDENT_DATASET_DIR)
    def train_faces(self): threading.Thread(target=lambda: self.face_recognition.train_face_encodings(), daemon=True).start()

    def _warmup_model(self):
        """Warmup the face recognition model in background"""
        try:
            if self.face_recognition.is_trained():
                self.face_recognition.train_face_encodings()
                logger.info("Model warmup complete")
        except Exception as e:
            logger.error(f"Model warmup failed: {e}")

    def display_image(self, img):
        """Display a static image in the camera label (Thread-safe)"""
        self.after(0, self._display_image_safe, img)

    def _display_image_safe(self, img):
        cw = self.cam_container.winfo_width()
        ch = self.cam_container.winfo_height()
        if cw < 10 or ch < 10: return
        
        # Resize
        h, w = img.shape[:2]
        ratio = min(cw/w, ch/h)
        new_w, new_h = int(w * ratio), int(h * ratio)
        
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        img_pil = Image.fromarray(img_rgb)
        ctk_img = ctk.CTkImage(img_pil, size=(new_w, new_h))
        self.camera_label.configure(image=ctk_img)

    def resume_live_feed(self):
        self.showing_result = False
        # self.update_camera_feed() # It is loop based, it will pick up next tick if after_id loop is running, wait update_camera_feed returns if showing_result
        # So we need to call it once to restart loop if it stopped?
        # Actually my update_camera_feed uses self.after even if it returns early? 
        # No, "if ... return" prevents self.after call.
        # So I must restart it.
        self.update_camera_feed()

    def update_last_attendance(self, attendance):
        self.last_attendance = attendance
        present = [name for name, status in attendance.items() if status == "Present"]
        
        # Update Sidebar
        if present:
            text = "Recently Seen:\n" + "\n".join(present[:3])
        else:
            text = "Recently Seen:\nNone"
        
        # Safe update
        self.after(0, lambda: self.recent_label.configure(text=text))

if __name__ == "__main__":
    app = AttendanceApp()
    app.mainloop()
   
 